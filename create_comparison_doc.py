# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_comparison_table():
    doc = Document()
    
    # 设置页面
    section = doc.sections[0]
    section.page_width = Inches(11.69)  # A4横向
    section.page_height = Inches(8.27)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    
    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('《成都工业职业技术学院外事工作管理办法》')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run('修订前后对照表')
    run2.bold = True
    run2.font.size = Pt(14)
    run2.font.color.rgb = RGBColor(192, 0, 0)
    
    # 说明
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run3 = info.add_run('说明：本表对比原制度（PDF）与修改后制度（Word）的内容差异，标注修改类型及备注。')
    run3.font.size = Pt(10)
    run3.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()  # 空白行
    
    # 创建表格
    headers = ['章节', '条款', '修改类型', '原制度内容', '修改后内容', '备注']
    col_widths = [Cm(2.5), Cm(1.8), Cm(2.2), Cm(6.5), Cm(6.5), Cm(4.5)]
    
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    
    # 表头
    header_row = table.rows[0]
    for i, (header, width) in enumerate(zip(headers, col_widths)):
        cell = header_row.cells[i]
        cell.width = width
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, 'D9E2F3')  # 浅蓝色背景
    
    # 差异数据
    differences = [
        {
            "chapter": "总则",
            "article": "第三条",
            "change_type": "文字调整",
            "original": "外籍人员聘请（用）管理、国（境）外援助项目管理",
            "modified": "外籍人员聘用（请）及学院教师国际培训等工作；学生处负责国际学生辅导员的统筹调配、国际学生的思想教育和日常管理工作",
            "note": "原条款中的\"外籍人员聘请（用）管理\"和\"国（境）外援助项目管理\"在修改版中拆分到各职能部门职责中表述"
        },
        {
            "chapter": "总则",
            "article": "第四条",
            "change_type": "新增引用",
            "original": "《成都市市级单位因公短期出国培训费用管理办法》",
            "modified": "《四川省因公临时出国经费管理办法》《成都市市级单位因公短期出国培训费用管理办法》",
            "note": "修改版新增了《四川省因公临时出国经费管理办法》作为经费管理依据"
        },
        {
            "chapter": "机构与职责",
            "article": "第五条",
            "change_type": "机构名称调整",
            "original": "领导小组办公室设在学院对外交流处",
            "modified": "领导小组办公室设在学院对外交流合作处",
            "note": "原版称\"对外交流处\"，修改版统一称为\"对外交流合作处\"（多处统一调整）"
        },
        {
            "chapter": "因公出国（境）管理",
            "article": "第十条",
            "change_type": "结构调整",
            "original": "1.院级领导出访按照党和国家的有关规定报批；\n2.各部门、二级学院处级领导干部出访，由学院外事工作领导小组商议，由院长办公会和党委会审议；\n3.各部门、二级学院其他人员出国（境）进行短期公务交流活动、短期学术交流、参加国际会议以及长期访学等活动，由院长办公会和党委会审议。",
            "modified": "所有因公出国（境）活动须按规定的程序分级、分类审批",
            "note": "修改版将审批层级条款编号直接写入正文，不再单独列出3个子项"
        },
        {
            "chapter": "因公出国（境）管理",
            "article": "第十一条",
            "change_type": "结构调整",
            "original": "有下列情形之一的出访项目和人员，一律不得申报和批准：\n1.目的不明确，项目无实质性内容的；\n2.出访国家、地区和机构与出访任务无直接关系的；\n……（共6项）",
            "modified": "有下列情形之一的出访项目和人员，一律不得申报和批准：[条款格式变化]",
            "note": "原版为明确列出的6项禁止情形，修改版中条款编号方式有所调整"
        },
        {
            "chapter": "因公出国（境）管理",
            "article": "第十四条",
            "change_type": "标点遗漏",
            "original": "所有因公出国（境）活动均须开展行前教育，出访人员须签订安全协议书和保密责任书。",
            "modified": "所有因公出国（境）活动均须开展行前教育出访人员须签订安全协议书和保密责任书。",
            "note": "修改版疑似遗漏了\"，\"（逗号），语句连接略显生硬"
        },
        {
            "chapter": "因公出国（境）管理",
            "article": "第十六条",
            "change_type": "标点符号",
            "original": "经组织部（人事处）、对外交流合作处审批，以下人员或情况可持因私证件出访：",
            "modified": "经组织部（人事处）对外交流合作处审批以下人员或情况可持因私证件出访",
            "note": "修改版缺少顿号\"、\"分隔，语句略显不通顺"
        },
        {
            "chapter": "学生出国（境）交流管理",
            "article": "第十九条",
            "change_type": "文字错误",
            "original": "学生出国（境）学习交流是指学院在读学生，通过学院与境外高校、教育机构、企业等合作项目出国（境）进行留学、游学、实习、社会实践等活动。",
            "modified": "学生出国（境）学习交流是指学院在读学生，通过学院与境外高校、教育机构、企业等合作项目出国（境进行留学、游学、实习、社会实践等活动。",
            "note": "修改版遗漏了\")\"，导致括号不匹配"
        },
        {
            "chapter": "签订国际合作协议管理",
            "article": "第二十八条",
            "change_type": "文字删减",
            "original": "经学院外事工作领导小组商议，院长办公会和党委会审议通过后，方可签订。",
            "modified": "经学院外事工作领导小组议，院长办公会和党委会审通过后，方可签订。",
            "note": "修改版疑似在排版过程中将\"商议\"缩略为\"议\"，\"审议\"缩略为\"审\"，疑似文字压缩/丢失"
        },
        {
            "chapter": "中外合作办学项目管理",
            "article": "第三十四条",
            "change_type": "数字错误",
            "original": "因教学要求提升，参与中外合作办学项目的语言类和专业类教师课时津贴可按正常课时津贴1.1倍计算。",
            "modified": "因教学要求提升，参与中外合作办学项目的语言类和专业类教师课时津贴可按正常课时津贴1. 倍计算。",
            "note": "修改版\"1.1\"变成了\"1. \"，小数点后数字丢失"
        },
        {
            "chapter": "国（境）外援助项目管理",
            "article": "第四十五条",
            "change_type": "文字遗漏",
            "original": "国（境）外援助项目指为改善学院教学、科研等办学条件，学院接受国（境）外友好捐赠或援助的资金、设备设施以及建设性项目。",
            "modified": "国（境）外援助项目指为改善学院教学、科研等办学条件，学院接受国（境）外友好捐赠或援助的资金设备设施以及建设性项目。",
            "note": "修改版遗漏了\"、\"（顿号），\"资金、设备设施\"变成\"资金设备设施\""
        },
        {
            "chapter": "外籍人员聘请（用）管理",
            "article": "第四十三条",
            "change_type": "标点/文字问题",
            "original": "组织部（人事处）是外籍聘请（用）人员的牵头部门，负责外籍聘请（用）人员的渠道建立、聘请（用）等工作……",
            "modified": "组织部（人事处）是外籍聘请（用）人员的牵头部门，负责外籍聘请（用）人员的渠道建立、聘请（用等工作……",
            "note": "修改版\"聘请（用）等工作\"少了\"等\"字后的停顿；\"谁使用、谁负责\"后多了空格"
        },
        {
            "chapter": "国际学生管理",
            "article": "第四十九条",
            "change_type": "文字错误",
            "original": "国际学生主要指以长（短）期进修、访学、实习等形式来学院接受学历教育或非学历教育的各类国（境）外人员。",
            "modified": "国际学生主要指以长（短）期进修、访学实习等形式来学院接受学历教育或非学历教育的各类国（境外人员。",
            "note": "1. 修改版遗漏了\"，\"（逗号）；2. \"国（境）外\"少了\")\"括号不匹配"
        },
        {
            "chapter": "整体结构",
            "article": "多处",
            "change_type": "格式调整",
            "original": "各条款以\"第X条\"格式编号，层级分明",
            "modified": "部分条款编号方式改变，有些直接嵌入正文",
            "note": "修改版在部分条款的呈现方式上有所调整，部分编号直接标注在正文中"
        },
    ]
    
    # 添加数据行
    for i, diff in enumerate(differences):
        row = table.add_row()
        
        # 交替颜色
        bg_color = 'F2F2F2' if i % 2 == 0 else 'FFFFFF'
        
        data = [diff['chapter'], diff['article'], diff['change_type'], 
                diff['original'], diff['modified'], diff['note']]
        
        for j, (cell, text) in enumerate(zip(row.cells, data)):
            cell.width = col_widths[j]
            p = cell.paragraphs[0]
            p.text = text
            p.runs[0].font.size = Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_shading(cell, bg_color)
            
            # 修改类型列居中
            if j == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].bold = True
                # 根据修改类型设置不同颜色
                if '错误' in diff['change_type'] or '遗漏' in diff['change_type'] or '删减' in diff['change_type']:
                    p.runs[0].font.color.rgb = RGBColor(192, 0, 0)  # 红色
                elif '新增' in diff['change_type']:
                    p.runs[0].font.color.rgb = RGBColor(0, 128, 0)  # 绿色
                else:
                    p.runs[0].font.color.rgb = RGBColor(0, 0, 128)  # 蓝色
    
    # 添加页脚说明
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_footer = footer.add_run(f'生成时间：{datetime.datetime.now().strftime("%Y年%m月%d日 %H:%M")}  |  共发现 {len(differences)} 处差异')
    run_footer.font.size = Pt(9)
    run_footer.font.color.rgb = RGBColor(128, 128, 128)
    
    # 保存
    output_path = r'C:\Users\成都工业学院\.qclaw\workspace\外事工作管理办法修订对照表.docx'
    doc.save(output_path)
    print(f'对照表已生成：{output_path}')
    return output_path

if __name__ == '__main__':
    create_comparison_table()
