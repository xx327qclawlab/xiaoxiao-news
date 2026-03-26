# -*- coding: utf-8 -*-
import os
import pdfplumber
from docx import Document

folder = r'C:\Users\成都工业学院\Desktop\新建文件夹 (4)'
files = os.listdir(folder)

pdf_file = None
docx_file = None

for f in files:
    if f.endswith('.pdf'):
        pdf_file = f
    elif f.endswith('.docx'):
        docx_file = f

print(f"PDF: {pdf_file}")
print(f"DOCX: {docx_file}")

# Extract PDF text
pdf_text = []
if pdf_file:
    pdf_path = os.path.join(folder, pdf_file)
    with pdfplumber.open(pdf_path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text:
                pdf_text.append(f"===== 第{i+1}页 =====\n{text}")
    
    with open(r'C:\Users\成都工业学院\.qclaw\workspace\pdf_content.txt', 'w', encoding='utf-8') as f:
        f.write('\n\n'.join(pdf_text))
    print(f"PDF提取完成，共 {len(pdf.pages)} 页")

# Extract DOCX text
docx_text = []
if docx_file:
    docx_path = os.path.join(folder, docx_file)
    doc = Document(docx_path)
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            docx_text.append(para.text)
    
    with open(r'C:\Users\成都工业学院\.qclaw\workspace\docx_content.txt', 'w', encoding='utf-8') as f:
        f.write('\n'.join(docx_text))
    print(f"DOCX提取完成，共 {len(doc.paragraphs)} 段")

print("\n内容已保存到文本文件")
